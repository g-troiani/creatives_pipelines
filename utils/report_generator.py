#
# UTILS/REPORT_GENERATOR.PY
#
 
import sqlite3
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
import os
import re
from ai_news_config import REPORT_DIRECTORY

def generate_daily_report(hacker_news_articles, techcrunch_articles):
    db_file = 'data/youtube_videos.db'
    connection = sqlite3.connect(db_file)
    cursor = connection.cursor()
    
    # Retrieve video data from the database
    query = "SELECT title, url, summary, channel FROM videos WHERE created_at >= date('now', '-1 day')"
    cursor.execute(query)
    video_data = cursor.fetchall()
    
    if video_data:
        # Create a new Word document
        document = Document()
        today = datetime.now().strftime("%Y-%m-%d")
        document.add_heading(f"AI Daily Brief - {today}", 0)
        
        # Add video summaries to the document
        document.add_heading("Video Summaries", 1)
        for i, video in enumerate(video_data, start=1):
            video_title, video_url, summary, video_channel = video
            
            # Format video title in bold
            paragraph = document.add_paragraph()
            run = paragraph.add_run(f"Video {i}: {video_title if video_title else 'N/A'}")
            run.bold = True
            run.font.size = Pt(14)
            
            # Add channel name
            paragraph = document.add_paragraph()
            run = paragraph.add_run(f"Channel: {video_channel if video_channel else 'N/A'}")
            run.font.size = Pt(12)
            
            # Add video URL as a hyperlink
            paragraph = document.add_paragraph()
            run = paragraph.add_run(f"URL: {video_url if video_url else 'N/A'}")
            run.font.size = Pt(12)
            run.underline = True
            run.font.color.rgb = None
            
            # Format the summary text
            if summary is not None:
                formatted_summary = re.sub(r'\*\*(.*?)\*\*', r'\1', summary)
                formatted_summary = re.sub(r'\n', r'\n', formatted_summary)
            else:
                formatted_summary = ""
            
            # Add summary
            paragraph = document.add_paragraph()
            lines = formatted_summary.split('\n')
            for line in lines:
                if line.startswith('Summary:'):
                    run = paragraph.add_run('Summary:')
                    run.bold = True
                    paragraph.add_run(line[8:])
                elif line.startswith('Key Message Bullet Points:'):
                    paragraph.add_run('\n')
                    run = paragraph.add_run('Key Message Bullet Points:')
                    run.bold = True
                elif line.startswith('Logical Grouping of Bullet Points:'):
                    paragraph.add_run('\n')
                    run = paragraph.add_run('Logical Grouping of Bullet Points:')
                    run.bold = True
                elif line.startswith('Expanded Details:'):
                    paragraph.add_run('\n')
                    run = paragraph.add_run('Expanded Details:')
                    run.bold = True
                else:
                    line = re.sub(r'\*(.*?)\*', r'\1', line)  # Remove single asterisks
                    paragraph.add_run(line)
                    paragraph.add_run('\n')
            
            # Add a page break between videos
            if i < len(video_data):
                document.add_page_break()
        
        # Add HackerNews section to the document
        document.add_heading("HackerNews", 1)
        for article in hacker_news_articles:
            # Add title
            paragraph = document.add_paragraph()
            run = paragraph.add_run(f"Title: {article['Title'] if article['Title'] else 'N/A'}")
            run.bold = True
            run.font.size = Pt(12)
            
            # Add URL link
            paragraph = document.add_paragraph()
            run = paragraph.add_run(f"URL: {article['Page URL'] if article['Page URL'] else 'N/A'}")
            run.bold = True
            run.underline = True
            run.font.color.rgb = None
            
            # Add summary
            paragraph = document.add_paragraph()
            run = paragraph.add_run("Summary:")
            run.bold = True
            paragraph.add_run(f" {article['Summary'] if article['Summary'] else 'N/A'}")
            
            # Add a page break between articles
            document.add_page_break()
        
        # Add TechCrunch section to the document
        document.add_heading("TechCrunch", 1)
        for article in techcrunch_articles:
            # Add title
            paragraph = document.add_paragraph()
            run = paragraph.add_run(f"Title: {article['Title'] if article['Title'] else 'N/A'}")
            run.bold = True
            run.font.size = Pt(12)
            
            # Add URL link
            paragraph = document.add_paragraph()
            run = paragraph.add_run(f"URL: {article['URL'] if article['URL'] else 'N/A'}")
            run.bold = True
            run.underline = True
            run.font.color.rgb = None
            
            # Add summary
            paragraph = document.add_paragraph()
            run = paragraph.add_run("Summary:")
            run.bold = True
            paragraph.add_run(f" {article['Summary'] if article['Summary'] else 'N/A'}")
            
            # Add a page break between articles
            document.add_page_break()
        
        # Save the document
        report_directory = REPORT_DIRECTORY
        os.makedirs(report_directory, exist_ok=True)
        report_path = os.path.join(report_directory, f"AI Daily Brief - {today}.docx")
        document.save(report_path)
        
        print(f"AI Daily Brief generated and saved at: {report_path}")
        return report_path
    else:
        print("No video data found for the last 24 hours.")
        return None
    
    cursor.close()
    connection.close()